"""Assemble draft/sections/*.md, resolving <equations/...>, <tables/...>, and
<figures/...> placement lines, plus draft/references/, into one super simple,
single-column, minimally formatted docx.

Usage: uv run python scripts/build_docx.py
Output: draft/manuscript/{YYMMDD}-{HHMM}.docx
"""
from __future__ import annotations

import csv
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

REPO = Path(__file__).resolve().parents[1]
SECTIONS = REPO / "draft/sections"
EQUATIONS = REPO / "draft/equations"
TABLES = REPO / "draft/tables"
FIGURES = REPO / "draft/figures"
REFERENCES = REPO / "draft/references"
OUT_DIR = REPO / "draft/manuscript"

SECTION_FILE_RE = re.compile(r"^\d+(\.\d+)*-.+\.md$")
PLACEMENT_RE = re.compile(r"^<(equations|tables|figures)/(.+)>$")

TEXT_RE = re.compile(r"\\text\{([^}]*)\}")
SUBSCRIPT_RE = re.compile(r"(_\{[^}]*\}|_[A-Za-z0-9])")
SYMBOLS = {
    r"\sum": "Σ", r"\cdot": "·", r"\in": "∈",
    r"\sigma": "σ", r"\lambda": "λ", r"\qquad": "   ",
}

# Front matter with no drafted content yet -- placeholders only, in paper order.
# Remove an entry here once its real draft/sections/*.md file exists.
FRONT_MATTER_PLACEHOLDERS = ["Authors", "Abstract", "Keywords", "Introduction"]
END_MATTER_PLACEHOLDERS = ["Conclusion"]


@dataclass
class Counters:
    figure: int = 0
    table: int = 0
    equation: int = 0


# Per-figure override width in inches, keyed by filename. Default is 4.5in.
FIGURE_WIDTH = {
    "2.2-hierarchy-diagram.png": 6.5,
    "2.3-training-curves.png": 6.5,
    "3.1-reliability-and-risk-coverage.png": 6.5,
    "3.2-gradcam-qualitative.png": 6.5,
}


def section_files():
    files = [f for f in SECTIONS.iterdir() if SECTION_FILE_RE.match(f.name)]
    return sorted(files, key=lambda f: tuple(int(p) for p in f.stem.split("-", 1)[0].split(".")))


def heading_level(stem: str) -> int:
    return stem.split("-", 1)[0].count(".") + 1


def heading_text(stem: str) -> str:
    num, _, slug = stem.partition("-")
    return f"{num} {slug.replace('-', ' ').capitalize()}"


def extract_braced(s: str, start: int) -> tuple[str, int]:
    depth = 0
    for i in range(start, len(s)):
        if s[i] == "{":
            depth += 1
        elif s[i] == "}":
            depth -= 1
            if depth == 0:
                return s[start + 1:i], i + 1
    raise ValueError(f"unbalanced braces in: {s}")


def replace_frac(text: str) -> str:
    # ponytail: hand-rolled, brace-depth-aware \frac{num}{den} -> (num) / (den).
    # Only covers what these three equations need; not a general LaTeX parser.
    out = []
    i = 0
    while i < len(text):
        if text[i:i + 5] == r"\frac":
            j = i + 5
            while j < len(text) and text[j] == " ":
                j += 1
            num, j = extract_braced(text, j)
            while j < len(text) and text[j] == " ":
                j += 1
            den, j = extract_braced(text, j)
            out.append(f"({replace_frac(num)}) / ({replace_frac(den)})")
            i = j
        else:
            out.append(text[i])
            i += 1
    return "".join(out)


def latexish_to_plain(text: str) -> str:
    text = TEXT_RE.sub(lambda m: m.group(1), text)
    text = replace_frac(text)
    for k, v in SYMBOLS.items():
        text = text.replace(k, v)
    return text


def add_equation_paragraph(doc: Document, counters: Counters, latex_line: str):
    text = latexish_to_plain(latex_line.strip())
    counters.equation += 1
    p = doc.add_paragraph()
    # Center tab centers the equation in the text column; right tab pins the
    # equation number to the margin -- the standard Word equation-numbering
    # layout, since paragraph-level CENTER alignment would center the number
    # along with the equation instead of pinning it.
    p.paragraph_format.tab_stops.add_tab_stop(Inches(3.25), WD_TAB_ALIGNMENT.CENTER)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    p.add_run("\t")
    for part in SUBSCRIPT_RE.split(text):
        if not part:
            continue
        is_sub = part.startswith("_")
        sub_text = part[2:-1] if part.startswith("_{") else part[1:] if is_sub else part
        run = p.add_run(sub_text)
        run.font.subscript = is_sub
        run.font.name = "Cambria Math"
        run.font.size = Pt(12)
    num_run = p.add_run(f"\t({counters.equation})")
    num_run.font.name = "Times New Roman"
    num_run.font.size = Pt(11)
    return counters.equation


def add_prose_paragraph(doc: Document, text: str):
    text = text.replace("\n", " ").strip()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Strip inline single-dollar math to plain text -- no special glyph handling,
    # keeps prose readable without a full inline-math renderer.
    text = re.sub(r"\$([^$]*)\$", lambda m: latexish_to_plain(m.group(1)), text)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


def add_caption(doc: Document, kind: str, number: int, asset_path: Path):
    caption_path = asset_path.parent / f"{asset_path.name}.caption.txt"
    text = caption_path.read_text(encoding="utf-8").strip() if caption_path.exists() else "[Caption to be added.]"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{kind} {number}. {text}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(9.5)


def _set_cell_borders(cell, **edges):
    # ponytail: minimal tcBorders writer, only the edges/attrs this script needs
    # (booktabs-style top/bottom rules, explicit "no vertical border" everywhere).
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for edge, spec in edges.items():
        tag = qn(f"w:{edge}")
        el = borders.find(tag)
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        for k, v in spec.items():
            el.set(qn(f"w:{k}"), str(v))


NO_BORDER = {"val": "nil"}
RULE = {"val": "single", "sz": "6", "color": "000000"}


def _style_table_borders(table):
    n_rows = len(table.rows)
    for r, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            edges = {"left": NO_BORDER, "right": NO_BORDER, "insideV": NO_BORDER}
            edges["top"] = RULE if r == 0 else NO_BORDER
            edges["bottom"] = RULE if r in (0, n_rows - 1) else NO_BORDER
            _set_cell_borders(cell, **edges)


def _to_number(text: str):
    try:
        return float(re.sub(r"[^0-9.\-]", "", text))
    except ValueError:
        return None


def _bold_cell(cell):
    for run in cell.paragraphs[0].runs:
        run.bold = True


# Data columns/rows that are sample sizes, not performance metrics -- no
# direction arrow, no bold-best comparison, even though they hold numbers.
NON_METRIC_LABELS = {"n boxes", "n", "condition", "metric", "finding", "group"}


def _direction(label: str):
    """Returns (clean_label, lower_is_better) or None if label isn't a metric."""
    if label.strip().lower() in NON_METRIC_LABELS:
        return None
    if "(lower is better)" in label.lower():
        return label.replace(" (lower is better)", "").replace(" (Lower is better)", ""), True
    return label, False


def _bold_best_values(table, header):
    # Only Table 1's shape (header[0] == "Metric") carries a per-row direction;
    # other row-wise tables (e.g. Table 2's Group/Finding/Flat/Hierarchical) name
    # one implied metric for the whole table, which defaults to higher-is-better.
    metric_labeled = bool(header) and header[0].strip().lower() == "metric"
    if "Flat" in header and "Hierarchical" in header:
        flat_i, hier_i = header.index("Flat"), header.index("Hierarchical")
        for row in table.rows[1:]:
            lower_better = False
            if metric_labeled:
                direction = _direction(row.cells[0].text)
                if direction is None:
                    continue
                lower_better = direction[1]
            a, b = _to_number(row.cells[flat_i].text), _to_number(row.cells[hier_i].text)
            if a is None or b is None or a == b:
                continue
            winner = flat_i if (a < b if lower_better else a > b) else hier_i
            _bold_cell(row.cells[winner])
    elif header and header[0] == "Condition":
        by_condition = {row.cells[0].text: row for row in table.rows[1:]}
        if "Flat" in by_condition and "Hierarchical" in by_condition:
            flat_row, hier_row = by_condition["Flat"], by_condition["Hierarchical"]
            for i in range(1, len(header)):
                if _direction(header[i]) is None:
                    continue
                a, b = _to_number(flat_row.cells[i].text), _to_number(hier_row.cells[i].text)
                if a is None or b is None or a == b:
                    continue
                _bold_cell(flat_row.cells[i] if a > b else hier_row.cells[i])


def _labels_with_arrows(labels):
    """Appends an up/down arrow to each label that names an actual metric
    (see NON_METRIC_LABELS for what's excluded, e.g. sample-size columns)."""
    out = []
    for label in labels:
        direction = _direction(label)
        if direction is None:
            out.append(label)
            continue
        clean, lower_better = direction
        out.append(f"{clean} {'↓' if lower_better else '↑'}")
    return out


def _set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for edge, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def add_table(doc: Document, counters: Counters, csv_path: Path):
    with open(csv_path, newline="", encoding="utf-8") as f:
        rows = list(csv.reader(f))
    if not rows:
        return
    header, body = rows[0], rows[1:]

    row_wise = "Flat" in header and "Hierarchical" in header
    column_wise = header and header[0] == "Condition"
    metric_labeled = bool(header) and header[0].strip().lower() == "metric"
    display_header = _labels_with_arrows(header) if column_wise else list(header)

    table = doc.add_table(rows=1, cols=len(header))
    for cell, text in zip(table.rows[0].cells, display_header):
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
    for row in body:
        cells = table.add_row().cells
        display_row = list(row)
        if metric_labeled:
            direction = _direction(row[0])
            if direction is not None:
                display_row[0] = f"{direction[0]} {'↓' if direction[1] else '↑'}"
        for i, (cell, text) in enumerate(zip(cells, display_row)):
            # Row-wise tables (Table 1, Table 2) center every column, including
            # the label column. Column-wise tables keep the identifier column left.
            cell.paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if (row_wise or i > 0) else WD_ALIGN_PARAGRAPH.LEFT
            )
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            run = cell.paragraphs[0].add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)
    _bold_best_values(table, header)
    _style_table_borders(table)
    for row in table.rows:
        for cell in row.cells:
            _set_cell_margins(cell)
    counters.table += 1
    add_caption(doc, "Table", counters.table, csv_path)
    doc.add_paragraph()


def add_figure(doc: Document, counters: Counters, png_path: Path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    width = FIGURE_WIDTH.get(png_path.name, 4.5)
    run.add_picture(str(png_path), width=Inches(width))
    counters.figure += 1
    add_caption(doc, "Figure", counters.figure, png_path)


def add_content_blocks(doc: Document, counters: Counters, text: str):
    blocks = re.split(r"\n\s*\n", text.strip())
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        placement = PLACEMENT_RE.match(block)
        if placement:
            kind, name = placement.groups()
            if kind == "equations":
                add_content_blocks(doc, counters, (EQUATIONS / name).read_text(encoding="utf-8"))
            elif kind == "tables":
                add_table(doc, counters, TABLES / name)
            elif kind == "figures":
                add_figure(doc, counters, FIGURES / name)
            continue
        if block.startswith("$$") and block.endswith("$$"):
            add_equation_paragraph(doc, counters, block.strip("$"))
            continue
        add_prose_paragraph(doc, block)


def add_placeholder_section(doc: Document, title: str):
    doc.add_heading(title, level=1)
    p = doc.add_paragraph()
    run = p.add_run(f"[{title} -- to be written.]")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


BIB_FIELD_RE = re.compile(r"(\w+)\s*=\s*\{([^{}]*)\}")
BIB_ENTRY_START_RE = re.compile(r"@(\w+)\s*\{([^,]+),")


def format_authors(raw: str) -> str:
    names = [n.strip() for n in raw.split(" and ") if n.strip()]
    if len(names) <= 1:
        return raw
    if len(names) == 2:
        return f"{names[0]} and {names[1]}"
    return ", ".join(names[:-1]) + f", and {names[-1]}"


def format_bib_entry(raw_entry: str) -> str:
    # ponytail: enough BibTeX parsing for our own consistently-formatted
    # entries (flat {value} fields, no nested braces) -- not a general parser.
    fields = {k.lower(): re.sub(r"\s+", " ", v).strip() for k, v in BIB_FIELD_RE.findall(raw_entry)}
    author = format_authors(fields.get("author", ""))
    title = fields.get("title", "")
    venue = fields.get("journal") or fields.get("booktitle") or fields.get("publisher") or ""
    year = fields.get("year", "")
    note = fields.get("note", "")
    parts = [author, f'"{title}"' if title else "", venue, year, note]
    return ", ".join(p for p in parts if p) + "."


def add_references(doc: Document):
    bib_files = list(REFERENCES.glob("*.bib"))
    entries = []
    if bib_files:
        raw = bib_files[0].read_text(encoding="utf-8")
        for raw_entry in re.split(r"\n(?=@)", raw):
            raw_entry = raw_entry.strip()
            if raw_entry:
                entries.append(format_bib_entry(raw_entry))

    doc.add_heading("References", level=1)
    if not entries:
        p = doc.add_paragraph()
        run = p.add_run("[References -- to be added.]")
        run.italic = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        return
    for i, entry in enumerate(entries, 1):
        p = doc.add_paragraph(f"[{i}] {entry}")
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    section = doc.sections[0]
    section.left_margin = section.right_margin = Inches(1)
    section.top_margin = section.bottom_margin = Inches(1)

    counters = Counters()
    files = section_files()
    if files and files[0].stem == "0-title":
        title_para = doc.add_paragraph(files[0].read_text(encoding="utf-8").strip())
        title_para.style = doc.styles["Title"]
        files = files[1:]
        for name in FRONT_MATTER_PLACEHOLDERS:
            add_placeholder_section(doc, name)

    for f in files:
        level = min(heading_level(f.stem), 3)
        doc.add_heading(heading_text(f.stem), level=level)
        add_content_blocks(doc, counters, f.read_text(encoding="utf-8"))

    for name in END_MATTER_PLACEHOLDERS:
        add_placeholder_section(doc, name)

    add_references(doc)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    now = datetime.now()
    out_path = OUT_DIR / f"{now:%y%m%d}-{now:%H%M}.docx"
    doc.save(out_path)
    print(f"wrote {out_path}")


if __name__ == "__main__":
    build()
